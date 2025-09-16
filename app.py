from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new document
doc = Document()

# Title
doc.add_heading("Dark Web Threat Intelligence Project", 0)

# Introduction
doc.add_heading("Introduction", level=1)
doc.add_paragraph(
    "This project focuses on detecting cybersecurity threats from the Dark Web. "
    "It collects simulated posts, analyzes potential risks, and generates alerts "
    "to help organizations proactively protect sensitive data."
)

# Tools and Languages
doc.add_heading("Tools & Languages Used", level=1)
doc.add_paragraph(
    "- Programming Language: Python\n"
    "- Libraries: BeautifulSoup, Requests, CSV\n"
    "- Workflow Automation: Power Automate (for alerting)\n"
    "- Visualization: Dashboard (custom / open-source tools)\n"
)

# Why Unique
doc.add_heading("Why This Project is Unique", level=1)
doc.add_paragraph(
    "Unlike traditional security systems that are mainly reactive, "
    "this project uses proactive intelligence gathering from the dark web. "
    "It identifies threats before they reach the organization, "
    "giving extra protection compared to conventional systems."
)

# Workflow Section with Image
doc.add_heading("Workflow of the System", level=1)
doc.add_paragraph(
    "The workflow involves scraping dark web data, analyzing text for threats, "
    "storing it in a dataset, and generating alerts if suspicious content is detected."
)
doc.add_picture("/mnt/data/A_flowchart_diagram_in_the_image_illustrates_a_cyb.png", width=Inches(5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Dashboard Section with Image
doc.add_heading("Dashboard Example", level=1)
doc.add_paragraph(
    "The dashboard shows detected threats, severity level, and alert messages. "
    "It provides security teams with real-time intelligence to act quickly."
)
doc.add_picture("/mnt/data/A_comparison_diagram_in_digital_format_contrasts_t.png", width=Inches(5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Threat Post Example with Image
doc.add_heading("Example of a Threat Post", level=1)
doc.add_paragraph(
    "Below is a simulated dark web threat post example. This demonstrates the type of "
    "data that can be captured, such as leaked credentials or credit card dumps."
)
doc.add_picture("/mnt/data/A_2D_digital_graphic_displays_four_cybersecurity_t.png", width=Inches(5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# How it Works
doc.add_heading("How It Works", level=1)
doc.add_paragraph(
    "1. Data Collection: Scrapes simulated dark web forums.\n"
    "2. Data Analysis: Filters for keywords (e.g., credentials, dumps, exploits).\n"
    "3. Threat Detection: Identifies suspicious posts.\n"
    "4. Alerting: Sends alert messages through Power Automate or email notifications.\n"
    "5. Dashboard: Displays threat levels and incidents in real-time."
)

# Save document
file_path = "/mnt/data/DarkWeb_Threat_Intel_Report_with_Images.docx"
doc.save(file_path)

file_path

